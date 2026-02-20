"""
scripts/build_docx.py
Reproducible Word document generator for the Hybrid Digital Twin manuscript.

Usage:
    python scripts/build_docx.py [--output manuscript/Hybrid_Digital_Twin_Seismic_RC.docx]

Steps:
    1. Concatenate manuscript/01_introduction.md → 07_acknowledgements.md
    2. Preprocess LaTeX: strip \\tag{N}, add equation numbers as plain text
    3. Run pandoc --mathml to convert to .docx (math rendered as OMML)
    4. Prepend title block (title, authors, affiliations) via python-docx
    5. Verify output: count OMML objects
"""

import argparse
import re
import shutil
import subprocess
import sys
import zipfile
from pathlib import Path

# ---------------------------------------------------------------------------
# Configuration
# ---------------------------------------------------------------------------

MANUSCRIPT_DIR = Path(__file__).parent.parent / "manuscript"
BUILD_DIR = Path(__file__).parent.parent / "_build"

SECTION_ORDER = [
    "01_introduction.md",
    "02_objectives.md",
    "03_methods.md",
    "04_results.md",
    "05_discussion.md",
    "06_conclusions.md",
    "07_acknowledgements.md",
]

TITLE_BLOCK = {
    "title": "Hybrid Digital Twin for Real-Time Seismic Response Prediction of Reinforced Concrete Buildings",
    "authors": "Miguel Rivera Ospina",
    "affiliation": "Department of Civil Engineering",
    "email": "[author email]",
    "year": "2026",
}

# ---------------------------------------------------------------------------
# Preprocessing
# ---------------------------------------------------------------------------


def preprocess_markdown(text: str) -> str:
    r"""Clean LaTeX tags and adjust equation numbering for pandoc.

    - Strips \tag{N} from display math (pandoc --mathml handles numbering poorly)
    - Keeps the equation number as a plain (N) immediately after the math block
    """
    lines = text.split("\n")
    output = []
    in_math = False
    pending_tag = None

    for line in lines:
        stripped = line.strip()

        # Track display math blocks $$ ... $$
        if stripped.startswith("$$") and not in_math:
            in_math = True
            # Extract \tag{N} if present on the opening line
            tag_match = re.search(r"\\tag\{(\d+)\}", line)
            if tag_match:
                pending_tag = tag_match.group(1)
                line = re.sub(r"\s*\\tag\{[^}]+\}", "", line)
            output.append(line)
            # Single-line $$ math $$ case
            if stripped.endswith("$$") and stripped != "$$":
                in_math = False
                if pending_tag:
                    output.append(f"\n*({pending_tag})*\n")
                    pending_tag = None
            continue

        if in_math:
            tag_match = re.search(r"\\tag\{(\d+)\}", line)
            if tag_match:
                pending_tag = tag_match.group(1)
                line = re.sub(r"\s*\\tag\{[^}]+\}", "", line)
            output.append(line)
            if stripped == "$$":
                in_math = False
                if pending_tag:
                    output.append(f"\n*({pending_tag})*\n")
                    pending_tag = None
            continue

        output.append(line)

    return "\n".join(output)


# ---------------------------------------------------------------------------
# Concatenation
# ---------------------------------------------------------------------------


def concatenate_sections(manuscript_dir: Path, sections: list[str]) -> str:
    """Concatenate markdown sections with section separators."""
    parts = []
    for filename in sections:
        path = manuscript_dir / filename
        if not path.exists():
            print(f"  [WARNING] Section not found: {path}", file=sys.stderr)
            continue
        content = path.read_text(encoding="utf-8")
        parts.append(content)

    return "\n\n---\n\n".join(parts)


# ---------------------------------------------------------------------------
# Title block (python-docx prepend)
# ---------------------------------------------------------------------------


def prepend_title_block(docx_path: Path, title_info: dict) -> None:
    """Insert title, authors, and affiliation paragraphs at the top of the docx."""
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    doc = Document(str(docx_path))

    # Build title paragraphs (inserted at position 0 in reverse order)
    inserts = [
        ("", 10, False, WD_ALIGN_PARAGRAPH.LEFT),  # spacer
        (
            f"© {title_info['year']} Authors. Open Access under CC BY 4.0.",
            8,
            False,
            WD_ALIGN_PARAGRAPH.CENTER,
        ),
        (
            f"* Corresponding Author: {title_info['email']}",
            9,
            False,
            WD_ALIGN_PARAGRAPH.CENTER,
        ),
        (title_info["affiliation"], 10, False, WD_ALIGN_PARAGRAPH.CENTER),
        (title_info["authors"], 11, True, WD_ALIGN_PARAGRAPH.CENTER),
        (title_info["title"], 14, True, WD_ALIGN_PARAGRAPH.CENTER),
    ]

    for _i, (text, size, bold, align) in enumerate(inserts):
        para = doc.add_paragraph()
        doc.paragraphs[0]._element.addprevious(para._element)
        para.alignment = align
        run = para.add_run(text)
        run.bold = bold
        run.font.size = Pt(size)

    doc.save(str(docx_path))
    print(f"  Title block prepended ({len(inserts)} paragraphs).")


# ---------------------------------------------------------------------------
# Verification
# ---------------------------------------------------------------------------


def verify_docx(docx_path: Path) -> dict:
    """Count OMML math objects and embedded content in the Word XML."""
    with zipfile.ZipFile(str(docx_path)) as z:
        doc_xml = z.read("word/document.xml").decode("utf-8", errors="replace")

    omml_count = doc_xml.count("<m:oMath>")
    table_count = doc_xml.count("<w:tbl>")
    fig_count = doc_xml.count("<w:drawing>") + doc_xml.count("<pic:pic")

    return {"omml_equations": omml_count, "tables": table_count, "figures": fig_count}


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------


def main() -> None:
    parser = argparse.ArgumentParser(description="Build Hybrid Digital Twin manuscript .docx")
    parser.add_argument(
        "--output",
        default=str(MANUSCRIPT_DIR / "Hybrid_Digital_Twin_Seismic_RC.docx"),
        help="Output .docx path",
    )
    parser.add_argument(
        "--no-title-block",
        action="store_true",
        help="Skip prepending the title block",
    )
    args = parser.parse_args()

    output_path = Path(args.output)
    BUILD_DIR.mkdir(exist_ok=True)

    # Step 1 — Concatenate
    print("Step 1: Concatenating manuscript sections...")
    raw_md = concatenate_sections(MANUSCRIPT_DIR, SECTION_ORDER)
    concat_path = BUILD_DIR / "concat.md"
    concat_path.write_text(raw_md, encoding="utf-8")
    print(f"  Concatenated {len(SECTION_ORDER)} sections → {concat_path}")

    # Step 2 — Preprocess
    print("Step 2: Preprocessing LaTeX tags...")
    processed_md = preprocess_markdown(raw_md)
    processed_path = BUILD_DIR / "processed.md"
    processed_path.write_text(processed_md, encoding="utf-8")
    print(f"  Preprocessed → {processed_path}")

    # Step 3 — Pandoc conversion
    print("Step 3: Running pandoc (--mathml → OMML)...")
    body_docx = BUILD_DIR / "body.docx"
    pandoc_cmd = [
        "pandoc",
        str(processed_path),
        "--mathml",
        "-s",
        "--from",
        "markdown+tex_math_dollars+raw_tex",
        "-o",
        str(body_docx),
    ]
    result = subprocess.run(pandoc_cmd, capture_output=True, text=True)
    if result.returncode != 0:
        print(f"  [ERROR] pandoc failed:\n{result.stderr}", file=sys.stderr)
        sys.exit(1)
    print(f"  pandoc completed → {body_docx} ({body_docx.stat().st_size // 1024} KB)")

    # Step 4 — Title block
    if not args.no_title_block:
        print("Step 4: Prepending title block...")
        prepend_title_block(body_docx, TITLE_BLOCK)
    else:
        print("Step 4: Skipping title block (--no-title-block).")

    # Step 5 — Copy to final destination
    print(f"Step 5: Copying to {output_path}...")
    shutil.copy2(str(body_docx), str(output_path))
    size_kb = output_path.stat().st_size // 1024
    print(f"  Output: {output_path} ({size_kb} KB)")

    # Step 6 — Verify
    print("Step 6: Verifying output...")
    stats = verify_docx(output_path)
    print(f"  OMML equations : {stats['omml_equations']}")
    print(f"  Tables         : {stats['tables']}")
    print(f"  Figures/draws  : {stats['figures']}")

    if stats["omml_equations"] < 5:
        print(
            "  [WARNING] Fewer than 5 OMML objects found — math may not have rendered correctly.",
            file=sys.stderr,
        )
    else:
        print(f"  [OK] Math rendering verified ({stats['omml_equations']} OMML objects).")

    print("\nDone.")


if __name__ == "__main__":
    main()
